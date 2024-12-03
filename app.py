import streamlit as st
import pdfplumber
import PyPDF2
import pandas as pd
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document
import re


def extract_metrics_from_text(full_text):
    patterns = {
        "HR": r"HR:\s*• Resting:\s*(\d+).*?• Deep Breathing:\s*(\d+-\d+).*?\(Δ:\s*(-?\d+)\s*\|\s*%Δ:\s*(-?\d+\.\d+%)\).*?• Valsalva:\s*(\d+-\d+).*?\(Δ:\s*(-?\d+)\s*\|\s*%Δ:\s*(-?\d+\.\d+%)\)",
        "BP": r"BP:\s*• Resting:\s*(\d+/\d+).*?• Deep Breathing:\s*(\d+/\d+).*?\(Δ:\s*(-?\d+)/(-?\d+)\s*\|\s*%Δ:\s*(-?\d+\.\d+%)\).*?• Valsalva:\s*(\d+/\d+).*?\(Δ:\s*(-?\d+)/(-?\d+)\s*\|\s*%Δ:\s*(-?\d+\.\d+%)\)",
        "RMF": r"RMF:\s*• Resting:\s*([\d.]+).*?• Deep Breathing:\s*([\d.]+).*?\(Δ:\s*(-?[\d.]+)\s*\|\s*%Δ:\s*(-?\d+\.\d+%)\).*?• Valsalva:\s*([\d.]+).*?\(Δ:\s*(-?[\d.]+)\s*\|\s*%Δ:\s*(-?\d+\.\d+%)\)",
        "LFa": r"LFa:\s*• Resting:\s*([\d.]+).*?• Deep Breathing:\s*([\d.]+).*?\(Δ:\s*([\d.]+)\s*\|\s*%Δ:\s*([\d+\.\d+%]+)\).*?• Valsalva:\s*([\d.]+).*?\(Δ:\s*([\d.]+)\s*\|\s*%Δ:\s*([\d+\.\d+%]+)\)",
        "HFa": r"HFa:\s*• Resting:\s*([\d.]+).*?• Deep Breathing:\s*([\d.]+).*?\(Δ:\s*([\d.]+)\s*\|\s*%Δ:\s*([\d+\.\d+%]+)\).*?• Valsalva:\s*([\d.]+).*?\(Δ:\s*([\d.]+)\s*\|\s*%Δ:\s*([\d+\.\d+%]+)\)",
        "LFa/HFa": r"LFa/HFa:\s*• Resting:\s*([\d.]+).*?• Deep Breathing:\s*([\d.]+).*?\(Δ:\s*(-?[\d.]+)\s*\|\s*%Δ:\s*(-?\d+\.\d+%)\).*?• Valsalva:\s*([\d.]+).*?\(Δ:\s*([\d.]+)\s*\|\s*%Δ:\s*([\d+\.\d+%]+)\)",
        "LF/HF": r"LF/HF:\s*• Resting:\s*([\d.]+).*?• Deep Breathing:\s*([\d.]+).*?\(Δ:\s*([\d.]+)\s*\|\s*%Δ:\s*([\d+\.\d+%]+)\).*?• Valsalva:\s*([\d.]+).*?\(Δ:\s*([\d.]+)\s*\|\s*%Δ:\s*([\d+\.\d+%]+)\)",
        "TSP": r"TSP:\s*• Resting:\s*(\d+).*?• Deep Breathing:\s*(\d+).*?\(Δ:\s*(\d+)\s*\|\s*%Δ:\s*(\d+\.\d+%)\).*?• Valsalva:\s*(\d+).*?\(Δ:\s*(\d+)\s*\|\s*%Δ:\s*(\d+\.\d+%)\)",
        "SDNN": r"SDNN:\s*• Resting:\s*(\d+).*?• Deep Breathing:\s*(\d+).*?\(Δ:\s*(\d+)\s*\|\s*%Δ:\s*(\d+\.\d+%)\).*?• Valsalva:\s*(\d+).*?\(Δ:\s*(\d+)\s*\|\s*%Δ:\s*(\d+\.\d+%)\)",
        "rmsSD": r"rmsSD:\s*• Resting:\s*(\d+).*?• Deep Breathing:\s*(\d+).*?\(Δ:\s*(\d+)\s*\|\s*%Δ:\s*(\d+\.\d+%)\).*?• Valsalva:\s*(\d+).*?\(Δ:\s*(\d+)\s*\|\s*%Δ:\s*(\d+\.\d+%)\)",
        "SpO2": r"SpO2:\s*• Resting:\s*(\d+).*?• Deep Breathing:\s*(\d+).*?\(Δ:\s*(\d+)\s*\|\s*%Δ:\s*(\d+\.\d+%)\).*?• Valsalva:\s*(\d+).*?\(Δ:\s*(\d+)\s*\|\s*%Δ:\s*(\d+\.\d+%)\)",
        "Other Metrics": r"(Autonomic Activity \(45-80 ms²\):\s*\d+\s*ms²|Valsalva Ratio \(>1\.2\):\s*[\d.]+|Functional Age \(<56 years\):\s*\d+\s*years)"
    }

    extracted_data = {}
    for metric, pattern in patterns.items():
        matches = re.findall(pattern, full_text, re.DOTALL)
        if matches:
            extracted_data[metric] = matches

    return extracted_data

    # Function to process and validate metrics
def process_extracted_data(extracted_data):
    processed_data = {}
    for metric, values in extracted_data.items():
        if values:
            if metric == "Other Metrics":
                processed_data[metric] = values
            else:
                processed_data[metric] = {
                    "Resting": values[0][0],
                    "Deep Breathing": values[0][1],
                    "Valsalva": values[0][4]
                }
    return processed_data


# Function to extract specific metrics from text
def extract_metrics_from_text(text):
    """
    Extract specific metrics like EEI, DDI, DEI, and AI from the page text.
    """
    patterns = {
        "EEI": r"EEI\s+([\d.]+)",
        "DDI": r"DDI\s+([\d.]+)",
        "DEI": r"DEI\s+([\d.]+)",
        "AI": r"AI\s+([\d.-]+)",
        "Right Ankle/Brachial Index": r"Right Ankle/Brachial Index:\s+([\d.]+)",
        "Right Toe/Brachial Index": r"Right Toe/Brachial Index:\s+([\d.]+)",
        "Left Ankle/Brachial Index": r"Left Ankle/Brachial Index:\s+([\d.]+)",
        "Left Toe/Brachial Index": r"Left Toe/Brachial Index:\s+([\d.]+)"
    }
    extracted_data = {}
    for key, pattern in patterns.items():
        match = re.search(pattern, text)
        extracted_data[key] = float(match.group(1)) if match else None
    return extracted_data

# Function to extract and validate data from a single page
def extract_and_validate_data(text):
    """
    Extracts and validates specific metrics against provided ranges.
    """
    metrics = {
        "Right Ankle/Brachial Index (1.0 - 1.4)": (1.0, 1.4),
        "Left Ankle/Brachial Index (1.0 - 1.4)": (1.0, 1.4),
        "Right Toe/Brachial Index (>0.75)": (0.75, None),
        "Left Toe/Brachial Index (>0.75)": (0.75, None),
        "EEI (0.3 - 0.7)": (0.3, 0.7),
        "DDI (0.3 - 0.7)": (0.3, 0.7),
        "DEI (0.3 - 0.7)": (0.3, 0.7),
        "AI (<-0.7)": (None, -0.7),
        "Reflection Index (0.65 - 0.85)": (0.65, 0.85),
        "Stiffness Index (<8.0 m/s)": (None, 8.0),
        "Cardiac Output (4.0 - 8.0 l/min)": (4.0, 8.0),
        "Mean Arterial Pressure (70 - 110 mmHg)": (70, 110),
        "C1 (>10.0 ml/mmHg)": (10.0, None),
        "C2 (>6.0 ml/mmHg)": (6.0, None),
        "Ventricular Extrasystole (<1)": (None, 1),
        "Atrial Extrasystole (<1)": (None, 1),
        "QRS (60 - 120 ms)": (60, 120),
        "QTc (350 - 460 ms)": (350, 460),
        "PR interval (120 - 200 ms)": (120, 200),
        "Body Mass Index (19 - 25)": (19, 25),
        "Stroke Volume (55 - 100 ml)": (55, 100),
        "Blood Volume (3 - 5 l)": (3, 5),
        "Artifacts (<1)": (None, 1),
        "ST seg (80 - 120)": (80, 120)
    }
    
    extracted_results = {}
    for metric, (low, high) in metrics.items():
        pattern = rf"{re.escape(metric.split('(')[0].strip())}.*?([\d.]+)"
        match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        if match:
            try:
                value = float(match.group(1))
                flag = (
                    (low is not None and value < low) or
                    (high is not None and value > high)
                )
                if flag:
                    extracted_results[metric] = f"{value} | **FLAG**"
                else:
                    extracted_results[metric] = f"{value}"
            except ValueError:
                extracted_results[metric] = "Invalid Value"
    return extracted_results

# Function to extract ECG Analysis data
def extract_ecg_data(pdf_path):
    reader = PdfReader(pdf_path)
    ecg_analysis_text = ""
    for page in reader.pages:
        text = page.extract_text()
        if "ECG ANALYSIS REPORT" in text:
            ecg_analysis_text = text
            break    

    if not ecg_analysis_text:
        print("ECG ANALYSIS REPORT section not found in the PDF.")
        return

    data_patterns = {
        "Heart Rate": {
            "Average Heart Rate": r"Average Heart Rate[:\s]+(\d+ bpm)",
            "Fastest rate": r"Fastest rate[:\s]+(\d+ bpm)",
            "Slowest rate": r"Slowest rate[:\s]+(\d+ bpm)",
            "Fastest minutely rate": r"Fastest minutely rate[:\s]+(\d+ bpm)",
            "Slowest minutely rate": r"Slowest minutely rate[:\s]+(\d+ bpm)"
        },
        
        "Ventricular Details": {
            "PVC - Ventricular Ectopy": r"PVC - Ventricular Ectopy[:\s]+(\d+ beats.*)",
            "Ventricular Couplet": r"Ventricular Couplet[:\s]+(\d+ episodes)"
        },

        "Pause / Block": {
            "Irregular / Artifact beat": r"Irregular / Artifact beat[:\s]+(\d+ beats.*)"
        },

        "Supraventricular Details": {
            "PAC - Supraventricular Ectopy": r"PAC - Supraventricular Ectopy[:\s]+(\d+ beats.*)",
            "Supraventricular Couplet": r"Supraventricular Couplet[:\s]+(\d+ episodes)"
        },
        
        "HRV Analysis": {
            "SDNN": r"SDNN[:\s]+(\d+ ms)",
            "SDNN Max": r"SDNN Max[:\s]+(\d+ ms)",
            "SDNN Min": r"SDNN Min[:\s]+(\d+ ms)"
        },
        
        "QRS Analysis": {
            "QRS": r"QRS[:\s]+(\d+ ms)",
            "QT / QTc": r"QT / QTc[:\s]+(\d+ ms / \d+ ms)",
            "QTc Max": r"QTc Max[:\s]+(\d+ ms)",
            "PR int / seg": r"PR int / seg[:\s]+(\d+ ms / \d+ ms)",
            "ST int / seg": r"ST int / seg[:\s]+(\d+ ms / \d+ ms)"
        },
    }
        

    extracted_data = {}
    for section, patterns in data_patterns.items():
        section_data = {}
        for key, pattern in patterns.items():
            match = re.search(pattern, ecg_analysis_text)
            section_data[key] = match.group(1).strip() if match else "N/A"
        extracted_data[section] = section_data

    return extracted_data

# Function to extract all data from the PDF
def extract_all_data(pdf_file):
    """
    Extract all relevant metrics (EEI, DDI, DEI, AI) and table data from the entire PDF.
    """
    metrics_list = []
    data_frames = []

    with pdfplumber.open(pdf_file) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            # Extract specific metrics from text
            page_text = page.extract_text()
            if page_text:
                metrics = extract_metrics_from_text(page_text)
                metrics_list.append(metrics)

            # Extract tables
            tables = page.extract_tables()
            for table in tables:
                if table:
                    headers = table[0]
                    unique_headers = [
                        f"Column_{i}" if header is None or headers.count(header) > 1 else header
                        for i, header in enumerate(headers)
                    ]
                    try:
                        df = pd.DataFrame(table[1:], columns=unique_headers)
                        data_frames.append(df)
                    except Exception as e:
                        st.warning(f"Error processing table on page {page_num}: {e}")

    # Combine all metrics into a single dictionary with the latest values
    combined_metrics = {key: None for key in [
        "EEI", "DDI", "DEI", "AI",
        "Right Ankle/Brachial Index", "Right Toe/Brachial Index",
        "Left Ankle/Brachial Index", "Left Toe/Brachial Index"
    ]}
    for metrics in metrics_list:
        for key, value in metrics.items():
            if value is not None:  # Update only if value is present
                combined_metrics[key] = value

    combined_df = pd.concat(data_frames, ignore_index=True) if data_frames else None
    return combined_df, combined_metrics

# Function to filter and process the required metrics from tables
def filter_required_metrics(df):
    """
    Filters and calculates only the required metrics for the output.
    """
    required_metrics = {
        "HR": "Heart Rate",
        "BP": "Blood Pressure",
        "RMF": "Respiratory Modulation Factor",
        "LFa": "Low-Frequency Activity",
        "HFa": "High-Frequency Activity",
        "LFa/HFa": "Low-Frequency/High-Frequency Ratio",
        "LF/HF": "LF/HF Ratio",
        "TSP": "Total Spectral Power",
        "SDNN": "Standard Deviation of NN Intervals",
        "rmsSD": "Root Mean Square of Successive Differences",
        "SpO2": "Oxygen Saturation",
    }
    formatted_data = {}
    for key, label in required_metrics.items():
        try:
            if key in df.columns:
                values = pd.to_numeric(df[key], errors="coerce").dropna().tolist()
                if len(values) >= 3:
                    resting, deep_breathing, valsalva = values[:3]
                    formatted_data[key] = {
                        "Resting": resting,
                        "Deep Breathing": deep_breathing,
                        "Valsalva": valsalva
                    }

        except Exception as e:
            st.warning(f"Error processing metric {label}: {e}")
    return formatted_data

# Function to create a Word document
def create_doc(metrics_data, extracted_metrics, validated_data, ecg_data):
    """
    Creates a Word document.
    """
    doc = Document()
    heading = doc.add_heading("ANS Report", level=1)
    run = heading.runs[0]
    run.font.size = Pt(26)  
    run.font.color.rgb = RGBColor(0, 0, 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add tabular metrics
    if metrics_data:
        for metric, values in metrics_data.items():
            doc.add_heading(metric, level=3)
            doc.add_paragraph(f"• Resting: {values['Resting']}")
            doc.add_paragraph(f"• Deep Breathing: {values['Deep Breathing']}")
            doc.add_paragraph(f"• Valsalva: {values['Valsalva']}")

    # Add extracted text metrics
    if extracted_metrics:
        heading1 = doc.add_heading("Cardiovascular Function", level=2)
        run1 = heading1.runs[0]
        run1.font.size = Pt(14)  
        run1.font.color.rgb = RGBColor(0, 0, 0)
        for metric, value in extracted_metrics.items():
            doc.add_paragraph(f"• {metric}: {value}")

    # Add validated metrics
    if validated_data:
        for metric, value in validated_data.items():
            doc.add_paragraph(f"• {metric}: {value}")


    if ecg_data:
        heading2 = doc.add_heading("ECG Analysis Report", level=2)
        run2 = heading2.runs[0]
        run2.font.size = Pt(16)  
        run2.font.color.rgb = RGBColor(0, 0, 0)
        for section, metrics in ecg_data.items():
            doc.add_heading(section, level=3)
            for key, value in metrics.items():
                doc.add_paragraph(f"• {key}: {value}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit app
def main():
    st.title("PDF Data Extractor")
    uploaded_file = st.file_uploader("Upload a PDF file", type="pdf")

    if uploaded_file is not None:
        st.write("Processing PDF...")
        try:
            # Extract all data from the PDF
            extracted_data, extracted_metrics = extract_all_data(uploaded_file)

            # Validate metrics from the first page
            with pdfplumber.open(uploaded_file) as pdf:
                first_page_text = pdf.pages[0].extract_text()
                validated_data = extract_and_validate_data(first_page_text)

            ecg_data = extract_ecg_data(uploaded_file)

            # Display extracted metrics
            if extracted_metrics:
                st.write("Extracted Metrics:")
                st.json(extracted_metrics)

            doc = Document()
            
            # Display validated metrics
            if validated_data:
                for metric, value in validated_data.items():
                    para = doc.add_paragraph()
                    run_metric = para.add_run(f"• {metric}: ")
                    run_metric.bold = False
                    if "**FLAG**" in value:
                        metric_value, _ = value.split("|")
                        run_value = para.add_run(metric_value.strip())
                        run_flag = para.add_run(" **FLAG**")
                        run_flag.bold = True
                    else:
                        run_value = para.add_run(value.strip())
                    run_value.bold = False

            # Filter and process only the required metrics
            if extracted_data is not None:
                metrics_data = filter_required_metrics(extracted_data)
                st.write("Processed Metrics from Tables:")
                st.json(metrics_data)

            # Generate Word document
            if st.button("Generate Word Document"):
                doc_buffer = create_doc(metrics_data, extracted_metrics, validated_data, ecg_data)
                st.download_button(
                    label="Download Summary Report",
                    data=doc_buffer,
                    file_name="metrics_summary.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
        except Exception as e:
            st.error(f"An error occurred: {e}")

if __name__ == "__main__":
    main()
