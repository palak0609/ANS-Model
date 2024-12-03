import streamlit as st
import pdfplumber
import re
from docx import Document
from io import BytesIO

# Function to extract data from the uploaded PDF
def extract_data_from_pdf(file):
    # Read PDF content
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text += page.extract_text()

    # Debugging: Print the raw text
    st.write("Extracted Raw Text:")
    st.text_area("Raw Text", text, height=300)

    # Define regex patterns
    metrics = {
        "HR": re.findall(
            r"HR.*?Resting: (\d+).*?Deep Breathing: (\d+)-(\d+).*?Valsalva: (\d+)-(\d+)",
            text,
            re.S,
        ),
        "BP": re.findall(
            r"BP.*?Resting: (\d+)/(\d+).*?Deep Breathing: (\d+)/(\d+).*?Valsalva: (\d+)/(\d+)",
            text,
            re.S,
        ),
        "SpO2": re.findall(
            r"SpO2.*?Resting: (\d+).*?Deep Breathing: (\d+).*?Valsalva: (\d+)", text, re.S
        ),
        "RMF": re.findall(
            r"RMF.*?Resting: ([\d.]+).*?Deep Breathing: ([\d.]+).*?Valsalva: ([\d.]+)",
            text,
            re.S,
        ),
        "LFa": re.findall(
            r"LFa.*?Resting: ([\d.]+).*?Deep Breathing: ([\d.]+).*?Valsalva: ([\d.]+)",
            text,
            re.S,
        ),
        "HFa": re.findall(
            r"HFa.*?Resting: ([\d.]+).*?Deep Breathing: ([\d.]+).*?Valsalva: ([\d.]+)",
            text,
            re.S,
        ),
    }

    # Debugging: Show extracted metrics
    st.write("Extracted Metrics:")
    st.write(metrics)

    return metrics

# Function to format the extracted metrics
def format_metrics(metrics):
    formatted = ""
    for key, values in metrics.items():
        if not values:  # Skip empty metrics
            continue
        formatted += f"{key}:\n"
        for value in values:
            if key == "HR":
                formatted += (
                    f"• Resting: {value[0]}\n"
                    f"• Deep Breathing: {value[1]}-{value[2]} (Δ: {int(value[2]) - int(value[1])} | %Δ: {((int(value[2]) - int(value[1])) / int(value[1]) * 100):.2f}%)\n"
                    f"• Valsalva: {value[3]}-{value[4]} (Δ: {int(value[4]) - int(value[3])} | %Δ: {((int(value[4]) - int(value[3])) / int(value[3]) * 100):.2f}%)\n"
                )
            elif key == "BP":
                formatted += (
                    f"• Resting: {value[0]}/{value[1]}\n"
                    f"• Deep Breathing: {value[2]}/{value[3]} (Δ: {int(value[2]) - int(value[0])}/{int(value[3]) - int(value[1])} | %Δ: {((int(value[2]) - int(value[0])) / int(value[0]) * 100):.2f}%/{((int(value[3]) - int(value[1])) / int(value[1]) * 100):.2f}%)\n"
                    f"• Valsalva: {value[4]}/{value[5]} (Δ: {int(value[4]) - int(value[0])}/{int(value[5]) - int(value[1])} | %Δ: {((int(value[4]) - int(value[0])) / int(value[0]) * 100):.2f}%/{((int(value[5]) - int(value[1])) / int(value[1]) * 100):.2f}%)\n"
                )
            elif key in ["SpO2", "RMF", "LFa", "HFa"]:
                formatted += (
                    f"• Resting: {value[0]}\n"
                    f"• Deep Breathing: {value[1]} (Δ: {float(value[1]) - float(value[0]):.2f} | %Δ: {((float(value[1]) - float(value[0])) / float(value[0]) * 100):.2f}%)\n"
                    f"• Valsalva: {value[2]} (Δ: {float(value[2]) - float(value[0]):.2f} | %Δ: {((float(value[2]) - float(value[0])) / float(value[0]) * 100):.2f}%)\n"
                )
    return formatted

# Function to create a Word document
def create_word_doc(formatted_data):
    doc = Document()
    doc.add_heading("Extracted Metrics Report", level=1)
    doc.add_paragraph(formatted_data)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit app
st.title("PDF Metrics Extractor")
uploaded_file = st.file_uploader("Upload your PDF file", type="pdf")

if uploaded_file:
    st.write("Processing your file...")
    metrics = extract_data_from_pdf(uploaded_file)
    formatted_data = format_metrics(metrics)

    # Display extracted data
    st.text_area("Extracted Data", formatted_data, height=300)

    # Generate Word document
    if formatted_data.strip():
        docx_file = create_word_doc(formatted_data)
        st.download_button(
            label="Download Word Document",
            data=docx_file,
            file_name="Extracted_Metrics.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        st.warning("No metrics found. Please check the uploaded PDF.")

